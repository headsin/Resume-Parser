import pdfplumber
import docx
import os
from io import BytesIO

def extract_text_from_file(file_source):
    """
    Extract text from PDF or DOCX.
    Handles both file paths and Streamlit UploadedFile objects.
    """
    text = ""

    # Detect if file_source is a Streamlit UploadedFile
    if hasattr(file_source, "name"):
        filename = file_source.name
        file_bytes = file_source.read()
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif ext == ".docx":
            doc = docx.Document(BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"

        else:
            raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

    else:
        # Local file path (for CLI/testing)
        _, ext = os.path.splitext(file_source)
        ext = ext.lower()

        if ext == ".pdf":
            with pdfplumber.open(file_source) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif ext == ".docx":
            doc = docx.Document(file_source)
            for para in doc.paragraphs:
                text += para.text + "\n"

        else:
            raise ValueError("Unsupported file type. Please use PDF or DOCX.")

    return text.strip()



